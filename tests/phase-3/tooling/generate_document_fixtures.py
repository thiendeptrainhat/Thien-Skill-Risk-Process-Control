"""Create synthetic DOCX/native-PDF/scan fixtures, not model test results.

Uses only installed workspace dependencies. Author via apply_patch; run into the
task fixture directories. The scan is drawn from scratch, not a real document.
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader

PRESET = {
    'name': 'standard_business_brief', 'header_pattern': 'memo_masthead',
    'page_inches': [8.5, 11], 'margin_inches': 1.0,
    'header_footer_inches': 0.492, 'width_dxa': 9360,
    'body_font': 'Calibri', 'body_size': 11, 'body_after': 6,
    'body_line': 1.10, 'table_widths_dxa': [1250, 3250, 2200, 2660],
    'named_overrides': {'title': '23pt, black, 0pt before, 4pt after; no border',
                        'table_text': '10pt to keep four factual columns readable'}
}

def xml_value(parent, name, values):
    node = parent.find(qn(name))
    if node is None:
        node = OxmlElement(name)
        parent.append(node)
    for key, value in values.items():
        node.set(qn(key), str(value))
    return node

def create_docx(data, path):
    doc = Document()
    doc.core_properties.author = 'Synthetic fixture generator'
    doc.core_properties.title = data['title']
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.492)
    tokens = {
      'Normal': (11, '000000', 0, 6, False),
      'Title': (23, '000000', 0, 4, True),
      'Subtitle': (11, '555555', 0, 6, False),
      'Heading 1': (16, '2E74B5', 16, 8, True),
      'Heading 2': (13, '2E74B5', 12, 6, True),
      'Heading 3': (12, '1F4D78', 8, 4, True),
      'Header': (9, '555555', 0, 0, False),
      'Footer': (9, '555555', 0, 0, False)
    }
    for name, (size, color, before, after, bold) in tokens.items():
        style = doc.styles[name]
        style.font.name, style.font.size = 'Calibri', Pt(size)
        style.font.bold, style.font.color.rgb = bold, RGBColor.from_string(color)
        style.font.italic = False
        if style.element.pPr is not None:
            for border in list(style.element.pPr.findall(qn('w:pBdr'))):
                style.element.pPr.remove(border)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
    sec.header.paragraphs[0].text = 'SYNTHETIC SOP - TEST FIXTURE'
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run('SOP-SYN-15 | ')
    field = OxmlElement('w:fldSimple')
    field.set(qn('w:instr'), 'PAGE')
    footer._p.append(field)
    doc.add_paragraph(data['title'], 'Title')
    doc.add_paragraph(data['subtitle'], 'Subtitle')
    doc.add_paragraph(data['authorization'])
    for section in data['sections']:
        doc.add_paragraph(section['heading'], 'Heading 2')
        for text in section.get('paragraphs', []):
            doc.add_paragraph(text)
        if 'table' not in section:
            continue
        rows = section['table']
        table = doc.add_table(rows=len(rows), cols=4)
        table.autofit = False
        pr = table._tbl.tblPr
        xml_value(pr, 'w:tblW', {'w:w':9360, 'w:type':'dxa'})
        xml_value(pr, 'w:tblInd', {'w:w':120, 'w:type':'dxa'})
        xml_value(pr, 'w:tblLayout', {'w:type':'fixed'})
        margins = xml_value(pr, 'w:tblCellMar', {})
        for side, width in [('top',80),('bottom',80),('start',120),('end',120)]:
            xml_value(margins, 'w:'+side, {'w:w':width,'w:type':'dxa'})
        borders = xml_value(pr, 'w:tblBorders', {})
        for side in ['top','left','bottom','right','insideH','insideV']:
            xml_value(borders,'w:'+side,{'w:val':'single','w:sz':4,'w:color':'B8BEC5'})
        for child in list(table._tbl.tblGrid):
            table._tbl.tblGrid.remove(child)
        for width in PRESET['table_widths_dxa']:
            col=OxmlElement('w:gridCol'); col.set(qn('w:w'),str(width)); table._tbl.tblGrid.append(col)
        for i,row in enumerate(rows):
            if i == 0:
                xml_value(table.rows[i]._tr.get_or_add_trPr(),'w:tblHeader',{})
            for j,text in enumerate(row):
                cell=table.cell(i,j); cell.text=text
                xml_value(cell._tc.get_or_add_tcPr(),'w:tcW',{'w:w':PRESET['table_widths_dxa'][j],'w:type':'dxa'})
                cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                if i == 0:
                    xml_value(cell._tc.get_or_add_tcPr(),'w:shd',{'w:fill':'F2F4F7'})
                for para in cell.paragraphs:
                    para.paragraph_format.space_after=Pt(3)
                    for run in para.runs:
                        run.font.size=Pt(10); run.bold=(i==0)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)

def create_pdf(data,path,font_dir):
    pdfmetrics.registerFont(TTFont('FixtureSans', str(font_dir/'DejaVuSans.ttf')))
    pdfmetrics.registerFont(TTFont('FixtureSansBold', str(font_dir/'DejaVuSans-Bold.ttf')))
    normal=ParagraphStyle('FixtureBody',fontName='FixtureSans',fontSize=10,leading=13,spaceAfter=6)
    title=ParagraphStyle('FixtureTitle',parent=normal,fontName='FixtureSansBold',fontSize=19,leading=23,spaceAfter=6)
    heading=ParagraphStyle('FixtureHeading',parent=normal,fontName='FixtureSansBold',fontSize=12,leading=15,spaceBefore=10,spaceAfter=6,textColor=colors.HexColor('#2E74B5'))
    cellstyle=ParagraphStyle('FixtureCell',parent=normal,fontSize=9,leading=12,spaceAfter=0)
    story=[Paragraph(data['title'],title),Paragraph(data['subtitle'],normal),Paragraph(data['authorization'],normal)]
    for section in data['sections']:
        story.append(Paragraph(section['heading'],heading))
        story.extend(Paragraph(t,normal) for t in section.get('paragraphs',[]))
        if 'table' in section:
            table=Table([[Paragraph(t,cellstyle) for t in row] for row in section['table']],colWidths=[w/20 for w in PRESET['table_widths_dxa']],repeatRows=1)
            table.setStyle(TableStyle([('GRID',(0,0),(-1,-1),0.4,colors.HexColor('#B8BEC5')),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#F2F4F7')),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(0,0),(-1,-1),6),('RIGHTPADDING',(0,0),(-1,-1),6),('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5)]))
            story.extend([table,Spacer(1,6)])
    path.parent.mkdir(parents=True,exist_ok=True)
    SimpleDocTemplate(str(path),pagesize=letter,leftMargin=72,rightMargin=72,topMargin=54,bottomMargin=54,title=data['title'],author='Synthetic fixture generator').build(story)

def create_scan(path,font_dir):
    image=Image.new('RGB',(1200,1600),'white'); draw=ImageDraw.Draw(image)
    font=ImageFont.truetype(str(font_dir/'DejaVuSans.ttf'),27)
    bold=ImageFont.truetype(str(font_dir/'DejaVuSans-Bold.ttf'),30)
    lines=[('SYNTHETIC SCAN - SOP-SYN-16',bold),('Dữ liệu tổng hợp; được phép dùng AI.',font),('Quy trình nhận hàng trả',bold),('Mục tiêu: đối chiếu hàng và hoàn tiền đúng.',font),('Nhận yêu cầu -> kiểm đếm -> duyệt -> hoàn tiền.',font)]
    y=90
    for text,face in lines:
        draw.text((80,y),text,fill='black',font=face); y+=70
    xs=[80,400,720,1120]; ys=[510,590,690,810]
    for x in xs: draw.line([(x,ys[0]),(x,ys[-1])],fill='black',width=2)
    for y in ys: draw.line([(xs[0],y),(xs[-1],y)],fill='black',width=2)
    values=[['Bước','Người thực hiện','Hồ sơ'],['Kiểm đếm','Nhân viên kho','Phiếu nhận'],['Duyệt hoàn tiền',None,'Ghi nhận duyệt']]
    small=ImageFont.truetype(str(font_dir/'DejaVuSans.ttf'),23)
    for i,row in enumerate(values):
        for j,text in enumerate(row):
            if text is not None: draw.text((xs[j]+14,ys[i]+20),text,fill='black',font=small)
    # Deliberate synthetic degradation: no recoverable authority value is encoded.
    draw.rectangle((417,708,702,789),fill=(218,218,218))
    for x in range(420,700,8): draw.line([(x,710),(x+2,785)],fill=(208,208,208),width=3)
    draw.text((80,910),'Ngoại lệ: giữ hàng để làm rõ chênh lệch.',fill='black',font=font)
    draw.text((80,970),'Lưu hồ sơ theo mã yêu cầu trả hàng.',fill='black',font=font)
    path.parent.mkdir(parents=True,exist_ok=True)
    c=canvas.Canvas(str(path),pagesize=letter,invariant=1)
    c.setTitle('Synthetic degraded scan SOP-SYN-16')
    c.drawImage(ImageReader(image),0,0,width=612,height=792)
    c.save()

def main():
    parser=argparse.ArgumentParser()
    parser.add_argument('--fixtures',type=Path,required=True)
    parser.add_argument('--font-dir',type=Path,required=True)
    args=parser.parse_args()
    data=json.loads((args.fixtures/'document-source.json').read_text())
    text=[data['title'],data['subtitle'],data['authorization']]
    for s in data['sections']:
        text.append(s['heading']); text.extend(s.get('paragraphs',[]))
        text.extend(' | '.join(row) for row in s.get('table',[]))
    chat=args.fixtures/'P1-U15-V01'/'input.md'; chat.parent.mkdir(parents=True,exist_ok=True)
    chat.write_text('\n\n'.join(text)+'\n')
    create_docx(data,args.fixtures/'P1-U15-V02'/'input.docx')
    create_pdf(data,args.fixtures/'P1-U15-V03'/'input.pdf',args.font_dir)
    create_scan(args.fixtures/'P1-U16-V01'/'input.pdf',args.font_dir)
    print(json.dumps({'created':['P1-U15-V01/input.md','P1-U15-V02/input.docx','P1-U15-V03/input.pdf','P1-U16-V01/input.pdf'],'preset':PRESET},ensure_ascii=False,indent=2))

if __name__=='__main__':
    main()
