#!/usr/bin/env python3
"""Read-only fixture semantic checks; not model behavior or visual review."""
from __future__ import annotations
import argparse
from datetime import datetime
import hashlib
import json
from pathlib import Path
import re
from zoneinfo import ZoneInfo
from docx import Document
from pypdf import PdfReader

def norm(value):
    return re.sub(r"\s+", " ", value).strip()

def digest(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()

def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--fixtures", type=Path, required=True)
    args = ap.parse_args()
    root = args.fixtures.resolve()
    source = json.loads((root / "document-source.json").read_text())
    units = [source[k] for k in ("title", "subtitle", "authorization")]
    for section in source["sections"]:
        units.append(section["heading"])
        units.extend(section.get("paragraphs", []))
        units.extend(cell for row in section.get("table", []) for cell in row)
    chat = root / "P1-U15-V01/input.md"
    word = root / "P1-U15-V02/input.docx"
    native = root / "P1-U15-V03/input.pdf"
    scan = root / "P1-U16-V01/input.pdf"
    doc = Document(word)
    word_text = "\n".join([p.text for p in doc.paragraphs] +
                          [c.text for t in doc.tables for r in t.rows for c in r.cells])
    pdf = PdfReader(native)
    scan_pdf = PdfReader(scan)
    data = {
        "chat": (chat, chat.read_text()),
        "docx": (word, word_text),
        "pdf_native": (native, "\n".join(p.extract_text() or "" for p in pdf.pages))
    }
    results = {}
    for key, (path, content) in data.items():
        missing = [i + 1 for i, text in enumerate(units) if norm(text) not in norm(content)]
        results[key] = {"path": str(path), "sha256": digest(path),
                        "content_units": len(units), "missing_units": missing,
                        "status": "pass" if not missing else "fail"}
    scan_text = "".join(p.extract_text() or "" for p in scan_pdf.pages)
    report = {
        "report_kind": "fixture_preflight_not_model_acceptance",
        "timestamp": datetime.now(ZoneInfo("Asia/Ho_Chi_Minh")).isoformat(),
        "semantic_parity": results,
        "docx_table_count": len(doc.tables), "native_pdf_page_count": len(pdf.pages),
        "scan": {"path": str(scan), "sha256": digest(scan), "pages": len(scan_pdf.pages),
                 "native_text_characters": len(scan_text),
                 "status": "pass" if not scan_text.strip() else "fail"},
        "status": "pass" if all(r["status"] == "pass" for r in results.values())
                  and len(doc.tables) == 1 and len(pdf.pages) == 1
                  and len(scan_pdf.pages) == 1 and not scan_text.strip() else "fail",
        "limitations": ["Checks fixture bytes/text units, not model extraction or answer parity.",
                        "Visual quality requires a separate actual render review.",
                        "Deliberately unreadable scan region has no hidden recoverable original value."]
    }
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0 if report["status"] == "pass" else 1

if __name__ == "__main__":
    raise SystemExit(main())
